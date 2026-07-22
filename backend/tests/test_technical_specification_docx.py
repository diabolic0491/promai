from datetime import date
from pathlib import Path

import pytest
from docx import Document

from app.services.technical_specification_docx import (
    MissingTemplateVariablesError,
    add_approval_date_variables,
    flatten_form_data,
    get_template_variables,
    render_docx_template,
)


def test_render_docx_replaces_split_and_nested_placeholders(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "generated.docx"
    document = Document()

    paragraph = document.add_paragraph()
    first_run = paragraph.add_run("Название: {{tz.")
    first_run.bold = True
    paragraph.add_run("title}}.")

    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = (
        "Заказчик: {{organization.full_name}}"
    )

    header = document.sections[0].header
    header.paragraphs[0].text = (
        "Дата: {{approval.date}}"
    )
    document.save(template_path)

    render_docx_template(
        template_path=template_path,
        output_path=output_path,
        values={
            "tz.title": "Тестовое ТЗ",
            "organization.full_name": (
                "ООО «Промас Инжиниринг»"
            ),
            "approval.date": "22.07.2026",
        },
        required_variables=[
            "{{tz.title}}",
            "organization.full_name",
        ],
    )

    generated = Document(output_path)

    assert generated.paragraphs[0].text == (
        "Название: Тестовое ТЗ."
    )
    assert generated.paragraphs[0].runs[0].bold is True
    assert generated.tables[0].cell(0, 0).text == (
        "Заказчик: ООО «Промас Инжиниринг»"
    )
    assert (
        generated.sections[0].header.paragraphs[0].text
        == "Дата: 22.07.2026"
    )
    assert get_template_variables(generated) == set()


def test_render_docx_reports_all_missing_variables(
    tmp_path: Path,
) -> None:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "generated.docx"
    document = Document()
    document.add_paragraph(
        "{{tz.title}} {{work.start_date}}"
    )
    document.save(template_path)

    with pytest.raises(
        MissingTemplateVariablesError
    ) as captured_error:
        render_docx_template(
            template_path=template_path,
            output_path=output_path,
            values={
                "tz.title": "Тестовое ТЗ",
            },
            required_variables=[
                "organization.director_position"
            ],
        )

    assert captured_error.value.variable_names == [
        "organization.director_position",
        "work.start_date",
    ]
    assert not output_path.exists()


def test_form_data_and_approval_date_variables() -> None:
    values = flatten_form_data(
        {
            "organization": {
                "director_position": "Директор",
            },
            "custom.code": "A-1",
        }
    )
    add_approval_date_variables(
        values,
        date(2026, 7, 22),
    )

    assert values == {
        "organization.director_position": "Директор",
        "custom.code": "A-1",
        "approval.date": "22.07.2026",
        "approval.day": "22",
        "approval.month": "июля",
        "approval.year": "2026",
    }
