from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from app.api.routes.contracts import (
    provide_contract_analysis_execution_context,
)
from app.main import app
from app.models.contract_analysis import (
    ContractAnalysisEvidenceReference as EvidenceModel,
)
from app.models.contract_analysis import (
    ContractAnalysisFinding as FindingModel,
)
from app.models.contract_analysis import (
    ContractAnalysisRun,
)
from app.models.user import User
from app.services import (
    contract_analysis_executor,
    contract_analysis_findings,
    contract_documents,
)

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


def build_contract_document() -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("ДОГОВОР ПОСТАВКИ")
    document.add_paragraph(
        "Оплата производится в течение 10 дней "
        "после поставки оборудования."
    )
    document.add_paragraph(
        "Поставщик несёт ответственность "
        "за просрочку."
    )
    document.save(stream)
    return stream.getvalue()


def build_contract_document_with_deadline_conflict(
) -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("ДОГОВОР ПОСТАВКИ")
    document.add_paragraph(
        "Поставщик обязуется выполнить "
        "предусмотренные настоящим Договором "
        "работы после поставки товара в течение "
        "5 рабочих дней после получения от "
        "Покупателя письменного уведомления о "
        "готовности к выполнению работ."
    )
    document.add_paragraph(
        "Срок выполнения предусмотренных "
        "настоящим Договором работ после поставки "
        "товара — не более 10 календарных дней "
        "после получения от Покупателя письменного "
        "уведомления о готовности к выполнению "
        "работ."
    )
    document.save(stream)
    return stream.getvalue()


def build_contract_document_with_payment_conflicts(
) -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_paragraph("ДОГОВОР ПОСТАВКИ")
    document.add_paragraph(
        "4.1. Покупатель уплачивает 30% стоимости "
        "товара в течение 3 банковских дней после "
        "подписания настоящего Договора."
    )
    document.add_paragraph(
        "4.2. Покупатель уплачивает оставшиеся "
        "80% стоимости товара в течение 5 "
        "банковских дней после поставки товара."
    )
    document.add_paragraph(
        "За просрочку оплаты Покупатель уплачивает "
        "Поставщику пеню в размере 0,10% от "
        "неоплаченной суммы за каждый день "
        "просрочки платежа."
    )
    document.add_paragraph(
        "При нарушении срока оплаты Покупатель "
        "уплачивает Поставщику пеню в размере "
        "0,15% от неоплаченной суммы за каждый "
        "день просрочки платежа."
    )
    document.save(stream)
    return stream.getvalue()


def create_contract_with_version(
    client: TestClient,
    *,
    unp: str,
    document_content: bytes | None = None,
) -> tuple[dict[str, object], dict[str, object]]:
    counterparty_response = client.post(
        "/counterparties",
        json={
            "unp": unp,
            "name": f"Контрагент {unp}",
            "short_name": f"Контрагент {unp}",
            "legal_address": "г. Минск",
        },
    )
    assert counterparty_response.status_code == 201
    counterparty = counterparty_response.json()
    contract_response = client.post(
        "/contracts",
        json={
            "counterparty_id": counterparty["id"],
            "number": f"ANALYSIS-{unp}",
            "title": "Договор для анализа",
            "contract_date": "2026-07-24",
            "currency": "BYN",
        },
    )
    assert contract_response.status_code == 201
    contract = contract_response.json()
    upload_response = client.post(
        (
            f"/contracts/{contract['id']}"
            "/versions/upload"
        ),
        files={
            "file": (
                "Договор.docx",
                (
                    document_content
                    if document_content is not None
                    else build_contract_document()
                ),
                DOCX_MEDIA_TYPE,
            ),
        },
    )
    assert upload_response.status_code == 201

    return contract, upload_response.json()


def build_policy():
    return (
        contract_analysis_findings
        .ContractAnalysisFindingsPolicy(
            policy_id="pilot-contract-policy",
            policy_version="2026-07-24",
            allowed_categories=(
                "payment",
                "liability",
                "delivery",
            ),
            allowed_severity_levels=(
                "medium",
                "high",
            ),
        )
    )


@dataclass
class PayloadExecutor:
    payload: object
    executor_name: str = "test_executor"
    model: str = "test-model"

    def execute(
        self,
        *,
        evidence_index,
        policy,
    ):
        del policy
        return (
            contract_analysis_executor
            .parse_executor_findings(
                payload=self.payload,
                evidence_index=evidence_index,
            )
        )


@dataclass
class FailingExecutor:
    executor_name: str = "test_executor"
    model: str = "test-model"

    def execute(
        self,
        *,
        evidence_index,
        policy,
    ):
        del evidence_index, policy
        raise (
            contract_analysis_executor
            .InvalidContractAnalysisExecutorResponseError(
                "RAW-PROVIDER-SECRET"
            )
        )


def configure_executor(payload: object) -> None:
    context = (
        contract_analysis_executor
        .ContractAnalysisExecutionContext(
            executor=PayloadExecutor(payload),
            policy=build_policy(),
        )
    )
    app.dependency_overrides[
        provide_contract_analysis_execution_context
    ] = lambda: context


@pytest.fixture
def storage_root(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> Path:
    monkeypatch.setattr(
        contract_documents.settings,
        "storage_root",
        str(tmp_path),
    )
    return tmp_path


def test_analysis_is_saved_and_returned_by_api(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
) -> None:
    del storage_root
    contract, version = (
        create_contract_with_version(
            client,
            unp="910000001",
        )
    )
    configure_executor(
        {
            "findings": [
                {
                    "category": "payment",
                    "severity_level": "medium",
                    "title": "Срок оплаты",
                    "description": (
                        "Срок оплаты требует проверки"
                    ),
                    "evidence": [
                        {
                            "block_id": None,
                            "quote": (
                                "Оплата производится "
                                "в течение 10 дней"
                            ),
                            "occurrence": 1,
                        }
                    ],
                }
            ]
        }
    )

    original_execute = (
        app.dependency_overrides[
            provide_contract_analysis_execution_context
        ]()
    )

    class ResolvingExecutor(PayloadExecutor):
        def execute(
            self,
            *,
            evidence_index,
            policy,
        ):
            block = next(
                block
                for block in evidence_index.blocks
                if "Оплата производится"
                in block.text
            )
            self.payload["findings"][0][
                "evidence"
            ][0]["block_id"] = block.block_id
            return super().execute(
                evidence_index=evidence_index,
                policy=policy,
            )

    context = (
        contract_analysis_executor
        .ContractAnalysisExecutionContext(
            executor=ResolvingExecutor(
                original_execute.executor.payload
            ),
            policy=build_policy(),
        )
    )
    app.dependency_overrides[
        provide_contract_analysis_execution_context
    ] = lambda: context

    response = client.post(

            f"/contracts/{contract['id']}"
            f"/versions/{version['version_number']}"
            "/analyses"

    )

    assert response.status_code == 201
    result = response.json()
    assert result["contract_id"] == contract["id"]
    assert result["document_version_id"] == version["id"]
    assert result["version_number"] == 1
    assert result["status"] == "completed"
    assert result["executor"] == "test_executor"
    assert result["model"] == "test-model"
    assert result["policy_id"] == (
        "pilot-contract-policy"
    )
    assert len(result["policy_sha256"]) == 64
    assert result["result_status"] == "machine_draft"
    assert result["requires_human_review"] is True
    assert result["error_code"] is None
    assert result["error_message"] is None
    assert len(result["findings"]) == 1
    finding = result["findings"][0]
    assert finding["ordinal"] == 1
    assert finding["category"] == "payment"
    assert finding["severity_level"] == "medium"
    evidence = finding["evidence_references"][0]
    assert evidence["ordinal"] == 1
    assert evidence["block_ordinal"] > 0
    assert evidence["quote"] == (
        "Оплата производится в течение 10 дней"
    )
    assert len(evidence["quote_sha256"]) == 64

    list_response = client.get(

            f"/contracts/{contract['id']}"
            "/versions/1/analyses"

    )
    assert list_response.status_code == 200
    assert len(list_response.json()) == 1
    assert "findings" not in list_response.json()[0]

    detail_response = client.get(

            f"/contracts/{contract['id']}"
            "/versions/1/analyses/"
            f"{result['id']}"

    )
    assert detail_response.status_code == 200
    detail_result = detail_response.json()
    assert detail_result["id"] == result["id"]
    assert detail_result["status"] == "completed"
    assert detail_result["result_id"] == (
        result["result_id"]
    )
    assert detail_result["findings"] == (
        result["findings"]
    )

    assert db_session.scalar(
        select(func.count())
        .select_from(ContractAnalysisRun)
        .where(
            ContractAnalysisRun.contract_id
            == contract["id"]
        )
    ) == 1
    assert db_session.scalar(
        select(func.count())
        .select_from(FindingModel)
        .where(
            FindingModel.analysis_run_id
            == result["id"]
        )
    ) == 1
    assert db_session.scalar(
        select(func.count())
        .select_from(EvidenceModel)
        .join(
            FindingModel,
            EvidenceModel.finding_id
            == FindingModel.id,
        )
        .where(
            FindingModel.analysis_run_id
            == result["id"]
        )
    ) == 1


def test_empty_findings_are_completed(
    client: TestClient,
    storage_root: Path,
) -> None:
    del storage_root
    contract, _version = create_contract_with_version(
        client,
        unp="910000002",
    )
    configure_executor({"findings": []})

    response = client.post(

            f"/contracts/{contract['id']}"
            "/versions/1/analyses"

    )

    assert response.status_code == 201
    result = response.json()
    assert result["status"] == "completed"
    assert result["result_status"] == "machine_draft"
    assert result["requires_human_review"] is True
    assert result["findings"] == []
    assert result["result_id"].startswith(
        "contract-findings-result-v1-"
    )


def test_deadline_conflict_is_saved_when_executor_misses_it(
    client: TestClient,
    storage_root: Path,
) -> None:
    del storage_root
    contract, _version = create_contract_with_version(
        client,
        unp="910000013",
        document_content=(
            build_contract_document_with_deadline_conflict()
        ),
    )
    configure_executor({"findings": []})

    response = client.post(
        (
            f"/contracts/{contract['id']}"
            "/versions/1/analyses"
        )
    )

    assert response.status_code == 201
    result = response.json()
    assert result["status"] == "completed"
    assert len(result["findings"]) == 1
    finding = result["findings"][0]
    assert finding["category"] == "delivery"
    assert finding["severity_level"] == "medium"
    assert finding["title"] == (
        "Несогласованность сроков выполнения работ"
    )
    assert len(
        finding["evidence_references"]
    ) == 2
    assert (
        "5 рабочих дней и 10 календарных дней"
        in finding["description"]
    )


def test_payment_conflicts_are_saved_when_executor_misses_them(
    client: TestClient,
    storage_root: Path,
) -> None:
    del storage_root
    contract, _version = create_contract_with_version(
        client,
        unp="910000014",
        document_content=(
            build_contract_document_with_payment_conflicts()
        ),
    )
    configure_executor({"findings": []})

    response = client.post(
        (
            f"/contracts/{contract['id']}"
            "/versions/1/analyses"
        )
    )

    assert response.status_code == 201
    result = response.json()
    assert result["status"] == "completed"
    assert len(result["findings"]) == 2
    assert tuple(
        finding["title"]
        for finding in result["findings"]
    ) == (
        "Несогласованность долей оплаты",
        "Несогласованность формулы расчёта пени",
    )
    assert all(
        finding["severity_level"] == "medium"
        for finding in result["findings"]
    )
    payment_finding = result["findings"][0]
    assert len(
        payment_finding["evidence_references"]
    ) == 2
    assert tuple(
        reference["quote"]
        for reference
        in payment_finding["evidence_references"]
    ) == (
        (
            "4.1. Покупатель уплачивает 30% "
            "стоимости товара в течение 3 "
            "банковских дней после подписания "
            "настоящего Договора."
        ),
        (
            "4.2. Покупатель уплачивает оставшиеся "
            "80% стоимости товара в течение 5 "
            "банковских дней после поставки товара."
        ),
    )


def test_semantically_unsupported_finding_is_not_saved(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
) -> None:
    del storage_root
    contract, _version = create_contract_with_version(
        client,
        unp="910000012",
    )
    configure_executor(
        {
            "findings": [
                {
                    "category": "payment",
                    "severity_level": "medium",
                    "title": "Отсутствие срока оплаты",
                    "description": (
                        "Срок оплаты не указан, "
                        "что создаёт риск спора."
                    ),
                    "evidence": [
                        {
                            "block_id": None,
                            "quote": (
                                "Оплата производится "
                                "в течение 10 дней"
                            ),
                            "occurrence": 1,
                        }
                    ],
                }
            ]
        }
    )
    original_context = (
        app.dependency_overrides[
            provide_contract_analysis_execution_context
        ]()
    )

    class ResolvingExecutor(PayloadExecutor):
        def execute(
            self,
            *,
            evidence_index,
            policy,
        ):
            block = next(
                block
                for block in evidence_index.blocks
                if "Оплата производится"
                in block.text
            )
            self.payload["findings"][0][
                "evidence"
            ][0]["block_id"] = block.block_id
            return super().execute(
                evidence_index=evidence_index,
                policy=policy,
            )

    context = (
        contract_analysis_executor
        .ContractAnalysisExecutionContext(
            executor=ResolvingExecutor(
                original_context.executor.payload
            ),
            policy=build_policy(),
        )
    )
    app.dependency_overrides[
        provide_contract_analysis_execution_context
    ] = lambda: context

    response = client.post(
        (
            f"/contracts/{contract['id']}"
            "/versions/1/analyses"
        )
    )

    assert response.status_code == 201
    result = response.json()
    assert result["status"] == "completed"
    assert result["findings"] == []
    assert db_session.scalar(
        select(func.count())
        .select_from(FindingModel)
        .where(
            FindingModel.analysis_run_id
            == result["id"]
        )
    ) == 0
    assert db_session.scalar(
        select(func.count())
        .select_from(EvidenceModel)
        .join(
            FindingModel,
            EvidenceModel.finding_id
            == FindingModel.id,
        )
        .where(
            FindingModel.analysis_run_id
            == result["id"]
        )
    ) == 0


def test_invalid_executor_result_is_saved_as_failed(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
) -> None:
    del storage_root
    contract, _version = create_contract_with_version(
        client,
        unp="910000003",
    )
    context = (
        contract_analysis_executor
        .ContractAnalysisExecutionContext(
            executor=FailingExecutor(),
            policy=build_policy(),
        )
    )
    app.dependency_overrides[
        provide_contract_analysis_execution_context
    ] = lambda: context

    response = client.post(

            f"/contracts/{contract['id']}"
            "/versions/1/analyses"

    )

    assert response.status_code == 502
    detail = response.json()["detail"]
    assert detail["status"] == "failed"
    assert detail["error_code"] == (
        "invalid_executor_response"
    )
    assert "RAW-PROVIDER-SECRET" not in (
        response.text
    )

    analysis_id = detail["analysis_id"]
    get_response = client.get(

            f"/contracts/{contract['id']}"
            "/versions/1/analyses/"
            f"{analysis_id}"

    )
    assert get_response.status_code == 200
    failed = get_response.json()
    assert failed["status"] == "failed"
    assert failed["result_id"] is None
    assert failed["result_status"] is None
    assert failed["requires_human_review"] is None
    assert failed["findings"] == []
    assert failed["error_code"] == (
        "invalid_executor_response"
    )
    assert "RAW-PROVIDER-SECRET" not in (
        failed["error_message"]
    )

    stored_run = db_session.get(
        ContractAnalysisRun,
        analysis_id,
    )
    assert stored_run is not None
    assert "RAW-PROVIDER-SECRET" not in (
        stored_run.error_message
    )


def test_disabled_analysis_returns_service_unavailable(
    client: TestClient,
    storage_root: Path,
) -> None:
    del storage_root
    contract, _version = create_contract_with_version(
        client,
        unp="910000004",
    )

    response = client.post(

            f"/contracts/{contract['id']}"
            "/versions/1/analyses"

    )

    assert response.status_code == 503
    assert response.json() == {
        "detail": "Анализ договоров выключен"
    }


def test_running_analysis_blocks_another_contract(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
) -> None:
    del storage_root
    running_contract, running_version = (
        create_contract_with_version(
            client,
            unp="910000005",
        )
    )
    queued_contract, _queued_version = (
        create_contract_with_version(
            client,
            unp="910000006",
        )
    )
    user_id = db_session.scalar(
        select(User.id).where(
            User.username == "test-admin"
        )
    )
    policy = build_policy()
    running = ContractAnalysisRun(
        contract_id=running_contract["id"],
        document_version_id=(
            running_version["id"]
        ),
        version_number=1,
        created_by_user_id=user_id,
        status="running",
        executor="test_executor",
        model="test-model",
        policy_id=policy.policy_id,
        policy_version=policy.policy_version,
        policy_sha256=(
            contract_analysis_findings
            .build_contract_analysis_policy_sha256(
                policy
            )
        ),
    )
    db_session.add(running)
    db_session.commit()
    configure_executor({"findings": []})

    response = client.post(

            f"/contracts/{queued_contract['id']}"
            "/versions/1/analyses"

    )

    assert response.status_code == 409
    assert response.json() == {
        "detail": (
            "Другой анализ договоров уже "
            "выполняется"
        )
    }


def test_running_analysis_returns_conflict(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
) -> None:
    del storage_root
    contract, version = create_contract_with_version(
        client,
        unp="910000007",
    )
    user_id = db_session.scalar(
        select(User.id).where(
            User.username == "test-admin"
        )
    )
    policy = build_policy()
    running = ContractAnalysisRun(
        contract_id=contract["id"],
        document_version_id=version["id"],
        version_number=1,
        created_by_user_id=user_id,
        status="running",
        executor="test_executor",
        model="test-model",
        policy_id=policy.policy_id,
        policy_version=policy.policy_version,
        policy_sha256=(
            contract_analysis_findings
            .build_contract_analysis_policy_sha256(
                policy
            )
        ),
    )
    db_session.add(running)
    db_session.commit()
    configure_executor({"findings": []})

    response = client.post(

            f"/contracts/{contract['id']}"
            "/versions/1/analyses"

    )

    assert response.status_code == 409
    assert response.json() == {
        "detail": (
            "Другой анализ договоров уже "
            "выполняется"
        )
    }


def test_analysis_routes_validate_parent_resources(
    client: TestClient,
) -> None:
    list_response = client.get(
        "/contracts/999999/versions/1/analyses"
    )
    assert list_response.status_code == 404
    assert list_response.json() == {
        "detail": "Договор не найден"
    }

    detail_response = client.get(

            "/contracts/999999/versions/1"
            "/analyses/999999"

    )
    assert detail_response.status_code == 404


def test_analysis_routes_require_authentication(
    anonymous_client: TestClient,
) -> None:
    response = anonymous_client.get(
        "/contracts/1/versions/1/analyses"
    )

    assert response.status_code == 401
    assert response.json() == {
        "detail": "Требуется аутентификация"
    }
