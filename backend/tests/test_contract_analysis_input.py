from hashlib import sha256
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from typing import Any

import pytest
from docx import Document

from app.services import contract_analysis_input
from app.services.contract_documents import (
    GeneratedContractFileNotFoundError,
)


def build_docx_content() -> bytes:
    stream = BytesIO()
    document = Document()
    document.sections[0].header.paragraphs[
        0
    ].text = "  PromAI\u00a0  header  "
    document.add_paragraph(
        "  1.   Предмет   договора  "
    )
    clause = document.add_paragraph(
        "1.1. Поставка"
    )
    clause.add_run().add_break()
    clause.add_run("оборудования")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Условие"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Цена | НДС"
    table.cell(1, 1).text = "1000 BYN"
    table.cell(1, 1).add_paragraph(
        "Срок: 5 дней"
    )
    document.sections[0].footer.paragraphs[
        0
    ].text = " Страница   1 "
    document.save(stream)
    return stream.getvalue()


def build_version(
    path: Path,
    content: bytes,
    *,
    file_sha256: str | None = None,
) -> SimpleNamespace:
    return SimpleNamespace(
        id=41,
        contract_id=17,
        version_number=3,
        file_name="Договор.docx",
        source="uploaded",
        storage_path=str(path),
        file_sha256=(
            file_sha256
            if file_sha256 is not None
            else sha256(content).hexdigest()
        ),
    )


def configure_version(
    monkeypatch: pytest.MonkeyPatch,
    *,
    version: SimpleNamespace,
    path: Path,
) -> list[dict[str, Any]]:
    calls: list[dict[str, Any]] = []

    def get_version(
        *,
        session: Any,
        contract_id: int,
        version_number: int,
    ) -> SimpleNamespace:
        calls.append(
            {
                "session": session,
                "contract_id": contract_id,
                "version_number": version_number,
            }
        )
        return version

    monkeypatch.setattr(
        contract_analysis_input,
        "get_contract_document_version",
        get_version,
    )
    monkeypatch.setattr(
        contract_analysis_input,
        "resolve_generated_file_path",
        lambda storage_path: path,
    )
    return calls


def test_prepare_contract_analysis_input_is_deterministic(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    content = build_docx_content()
    path = tmp_path / "contract.docx"
    path.write_bytes(content)
    version = build_version(path, content)
    calls = configure_version(
        monkeypatch,
        version=version,
        path=path,
    )
    session = object()
    expected_text = """[HEADER 1]
PromAI header

[BODY]
1. Предмет договора

1.1. Поставка
оборудования

[TABLE]
| Условие | Значение |
| Цена \\| НДС | 1000 BYN\\n\\nСрок: 5 дней |
[/TABLE]

[FOOTER 1]
Страница 1"""

    first_result = (
        contract_analysis_input
        .prepare_contract_analysis_input(
            session,
            contract_id=17,
            version_number=3,
        )
    )
    second_result = (
        contract_analysis_input
        .prepare_contract_analysis_input(
            session,
            contract_id=17,
            version_number=3,
        )
    )

    assert first_result == second_result
    assert first_result.contract_id == 17
    assert first_result.document_version_id == 41
    assert first_result.version_number == 3
    assert first_result.file_name == "Договор.docx"
    assert first_result.source == "uploaded"
    assert first_result.source_file_sha256 == (
        sha256(content).hexdigest()
    )
    assert first_result.text == expected_text
    encoded_text = expected_text.encode("utf-8")
    assert first_result.extracted_text_sha256 == (
        sha256(encoded_text).hexdigest()
    )
    assert first_result.source_file_size_bytes == len(
        content
    )
    assert first_result.extracted_text_characters == len(
        expected_text
    )
    assert first_result.extracted_text_size_bytes == len(
        encoded_text
    )
    assert calls == [
        {
            "session": session,
            "contract_id": 17,
            "version_number": 3,
        },
        {
            "session": session,
            "contract_id": 17,
            "version_number": 3,
        },
    ]


def test_prepare_rejects_replaced_document(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    content = build_docx_content()
    path = tmp_path / "contract.docx"
    path.write_bytes(content)
    version = build_version(
        path,
        content,
        file_sha256=sha256(
            b"original content"
        ).hexdigest(),
    )
    configure_version(
        monkeypatch,
        version=version,
        path=path,
    )

    with pytest.raises(
        contract_analysis_input
        .ContractAnalysisDocumentIntegrityError
    ):
        (
            contract_analysis_input
            .prepare_contract_analysis_input(
                object(),
                contract_id=17,
                version_number=3,
            )
        )


@pytest.mark.parametrize(
    "stored_hash",
    [
        "",
        "not-a-valid-sha256",
    ],
)
def test_prepare_rejects_missing_or_invalid_stored_hash(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
    stored_hash: str,
) -> None:
    content = build_docx_content()
    path = tmp_path / "contract.docx"
    path.write_bytes(content)
    version = build_version(
        path,
        content,
        file_sha256=stored_hash,
    )
    configure_version(
        monkeypatch,
        version=version,
        path=path,
    )

    with pytest.raises(
        contract_analysis_input
        .ContractAnalysisDocumentIntegrityError
    ):
        (
            contract_analysis_input
            .prepare_contract_analysis_input(
                object(),
                contract_id=17,
                version_number=3,
            )
        )


def test_prepare_rejects_invalid_docx(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    content = b"not a docx archive"
    path = tmp_path / "contract.docx"
    path.write_bytes(content)
    version = build_version(path, content)
    configure_version(
        monkeypatch,
        version=version,
        path=path,
    )

    with pytest.raises(
        contract_analysis_input
        .InvalidContractAnalysisDocumentError
    ):
        (
            contract_analysis_input
            .prepare_contract_analysis_input(
                object(),
                contract_id=17,
                version_number=3,
            )
        )


def test_prepare_rejects_empty_docx(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    stream = BytesIO()
    Document().save(stream)
    content = stream.getvalue()
    path = tmp_path / "contract.docx"
    path.write_bytes(content)
    version = build_version(path, content)
    configure_version(
        monkeypatch,
        version=version,
        path=path,
    )

    with pytest.raises(
        contract_analysis_input
        .EmptyContractAnalysisDocumentError
    ):
        (
            contract_analysis_input
            .prepare_contract_analysis_input(
                object(),
                contract_id=17,
                version_number=3,
            )
        )


def test_prepare_rejects_unavailable_document(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    content = build_docx_content()
    path = tmp_path / "missing.docx"
    version = build_version(path, content)
    monkeypatch.setattr(
        contract_analysis_input,
        "get_contract_document_version",
        lambda **kwargs: version,
    )

    def reject_path(
        storage_path: str,
    ) -> Path:
        raise GeneratedContractFileNotFoundError

    monkeypatch.setattr(
        contract_analysis_input,
        "resolve_generated_file_path",
        reject_path,
    )

    with pytest.raises(
        contract_analysis_input
        .ContractAnalysisDocumentUnavailableError
    ):
        (
            contract_analysis_input
            .prepare_contract_analysis_input(
                object(),
                contract_id=17,
                version_number=3,
            )
        )
