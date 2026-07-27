import json
from hashlib import sha256
from io import BytesIO
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy.orm import Session

from app.models.contract import Contract
from app.models.contract_document_version import (
    ContractDocumentVersion,
)
from app.models.document_template import (
    DocumentTemplate,
)
from app.services import (
    contract_documents,
    document_templates,
)
from app.services.technical_specification_docx import (
    get_template_variables,
    iter_document_paragraphs,
)


DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)


@pytest.fixture
def storage_root(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> Path:
    monkeypatch.setattr(
        contract_documents.settings,
        "storage_root",
        str(tmp_path),
    )
    monkeypatch.setattr(
        document_templates.settings,
        "storage_root",
        str(tmp_path),
    )
    return tmp_path


def create_counterparty(
    client: TestClient,
) -> dict[str, Any]:
    response = client.post(
        "/counterparties",
        json={
            "unp": "900000101",
            "name": "ООО «Покупатель»",
            "short_name": "ООО «Покупатель»",
            "legal_address": "г. Минск",
        },
    )

    assert response.status_code == 201
    return response.json()


def update_organization_profile(
    client: TestClient,
) -> dict[str, Any]:
    response = client.patch(
        "/organization-profile",
        json={
            "name": "ООО «Промас Инжиниринг»",
            "short_name": "ООО «Промас»",
            "unp": "590000001",
            "legal_address": "г. Гродно",
            "email": "info@example.by",
            "phone": "+375 00 000-00-00",
            "director_name": "Иванов И.И.",
            "director_position": "Директор",
            "bank_name": "ОАО «Тест Банк»",
            "bank_account": "BY00TEST00000000000000000000",
            "bank_code": "TESTBY2X",
        },
    )

    assert response.status_code == 200
    return response.json()


def create_template_file(
    path: Path,
    *,
    missing_variables_only: bool = False,
) -> None:
    document = Document()

    if missing_variables_only:
        document.add_paragraph(
            "{{contract.number}} "
            "{{contract.subject}} "
            "{{counterparty.account}}"
        )
    else:
        paragraph = document.add_paragraph(
            "Договор № "
        )
        first_run = paragraph.add_run("{{contract.")
        first_run.bold = True
        paragraph.add_run("number}} от ")
        paragraph.add_run(
            "{{contract.day}} "
            "{{contract.month}} "
            "{{contract.year}}"
        )

        document.add_paragraph(
            "Предмет: {{contract.subject}}; "
            "сумма: {{contract.amount}} {{contract.currency}}"
        )
        document.add_paragraph(
            "Поставщик: {{organization.name}}, "
            "{{organization.address}}, "
            "{{organization.account}}, "
            "{{organization.authority}}"
        )

        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = (
            "Покупатель: {{counterparty.name}}, "
            "{{counterparty.address}}, "
            "{{counterparty.account}}, "
            "{{counterparty.director_name}}"
        )

    document.save(path)


def upload_template(
    client: TestClient,
    template_path: Path,
    *,
    template_type: str = "contract",
    required_variables: list[str] | None = None,
) -> dict[str, Any]:
    with template_path.open("rb") as template_file:
        response = client.post(
            "/document-templates",
            data={
                "name": "Шаблон договора",
                "template_type": template_type,
                "required_variables": json.dumps(
                    required_variables or [],
                    ensure_ascii=False,
                ),
            },
            files={
                "file": (
                    "contract-template.docx",
                    template_file,
                    DOCX_MEDIA_TYPE,
                ),
            },
        )

    assert response.status_code == 201
    return response.json()


def contract_form_data() -> dict[str, Any]:
    return {
        "contract": {
            "city": "Гродно",
            "subject": "Поставка оборудования",
            "delivery_scope": "Поставка и монтаж",
            "amount_words": "Одна тысяча рублей",
            "vat_amount": "166.67",
            "vat_amount_words": (
                "Сто шестьдесят шесть рублей 67 копеек"
            ),
            "payment_terms": "Оплата в течение 10 дней",
            "delivery_period": "30 календарных дней",
            "delivery_address": "г. Минск",
            "number": "НЕ ДОЛЖНО ПОПАСТЬ В ДОКУМЕНТ",
        },
        "organization": {
            "authority": "Устава",
            "name": "НЕ ДОЛЖНО ПОПАСТЬ В ДОКУМЕНТ",
        },
        "counterparty": {
            "account": "BY11TEST11111111111111111111",
            "bank": "ЗАО «Банк Покупателя»",
            "bic": "BUYERBY2X",
            "director_name": "Петров П.П.",
            "director_position": "Директор",
            "authority": "Устава",
            "name": "НЕ ДОЛЖНО ПОПАСТЬ В ДОКУМЕНТ",
        },
    }


def create_contract(
    client: TestClient,
    counterparty_id: int,
    *,
    template_id: int | None,
    form_data: dict[str, Any] | None = None,
) -> dict[str, Any]:
    payload: dict[str, Any] = {
        "counterparty_id": counterparty_id,
        "template_id": template_id,
        "number": "Д-101/26",
        "title": "Договор поставки",
        "contract_date": "2026-07-22",
        "start_date": "2026-07-22",
        "end_date": "2026-12-31",
        "amount": "1000.00",
        "currency": "BYN",
        "notes": "Тестовая генерация",
        "owner_role": "supplier",
        "counterparty_role": "buyer",
        "form_data": form_data or {},
    }
    response = client.post(
        "/contracts",
        json=payload,
    )

    assert response.status_code == 201
    return response.json()


def build_contract_document_content(
    text: str = "Внешний договор контрагента",
) -> bytes:
    stream = BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(stream)
    return stream.getvalue()


def document_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    return "\n".join(
        paragraph.text
        for paragraph in iter_document_paragraphs(
            document
        )
    )


def test_contract_template_can_be_uploaded(
    client: TestClient,
    storage_root: Path,
) -> None:
    template_path = storage_root / "source.docx"
    create_template_file(template_path)

    template = upload_template(
        client,
        template_path,
    )

    assert template["template_type"] == "contract"
    assert template["file_name"] == (
        "contract-template.docx"
    )
    assert template["required_variables"] == [
        "contract.amount",
        "contract.currency",
        "contract.day",
        "contract.month",
        "contract.number",
        "contract.subject",
        "contract.year",
        "counterparty.account",
        "counterparty.address",
        "counterparty.director_name",
        "counterparty.name",
        "organization.account",
        "organization.address",
        "organization.authority",
        "organization.name",
    ]


def test_template_upload_combines_and_normalizes_variables(
    client: TestClient,
    storage_root: Path,
) -> None:
    template_path = storage_root / "variables.docx"
    document = Document()
    document.add_paragraph(
        "Номер: {{contract.number}}"
    )
    document.sections[0].header.paragraphs[0].text = (
        "Дата: {{approval.date}}"
    )
    document.save(template_path)

    template = upload_template(
        client,
        template_path,
        required_variables=[
            " {{ organization.authority }} ",
            "contract.number",
            "organization.authority",
        ],
    )

    assert template["required_variables"] == [
        "approval.date",
        "contract.number",
        "organization.authority",
    ]


def test_template_upload_rejects_invalid_required_variable(
    client: TestClient,
    storage_root: Path,
) -> None:
    template_path = storage_root / "invalid-required.docx"
    create_template_file(template_path)

    with template_path.open("rb") as template_file:
        response = client.post(
            "/document-templates",
            data={
                "name": "Некорректный шаблон",
                "template_type": "contract",
                "required_variables": json.dumps(
                    ["contract subject"],
                    ensure_ascii=False,
                ),
            },
            files={
                "file": (
                    "invalid-required.docx",
                    template_file,
                    DOCX_MEDIA_TYPE,
                ),
            },
        )

    assert response.status_code == 422
    assert response.json() == {
        "detail": (
            "required_variables должен быть "
            "JSON-массивом корректных имён переменных"
        ),
    }
    templates_directory = storage_root / "templates"
    assert (
        not templates_directory.exists()
        or not list(templates_directory.iterdir())
    )


def test_template_upload_rejects_invalid_docx_variable(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
) -> None:
    template_path = storage_root / "invalid-docx.docx"
    document = Document()
    document.add_paragraph(
        "Предмет: {{contract subject}}"
    )
    document.save(template_path)

    templates_count_before = (
        db_session.query(DocumentTemplate).count()
    )

    with template_path.open("rb") as template_file:
        response = client.post(
            "/document-templates",
            data={
                "name": "Некорректный шаблон",
                "template_type": "contract",
                "required_variables": "[]",
            },
            files={
                "file": (
                    "invalid-docx.docx",
                    template_file,
                    DOCX_MEDIA_TYPE,
                ),
            },
        )

    assert response.status_code == 422
    assert response.json() == {
        "detail": {
            "message": (
                "DOCX содержит некорректные "
                "имена переменных"
            ),
            "invalid_variables": [
                "contract subject",
            ],
        },
    }
    assert not list(
        (storage_root / "templates").iterdir()
    )
    assert (
        db_session.query(DocumentTemplate).count()
        == templates_count_before
    )


def test_generate_and_download_contract_docx(
    client: TestClient,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    update_organization_profile(client)
    template_path = storage_root / "source.docx"
    create_template_file(template_path)
    template = upload_template(
        client,
        template_path,
    )
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=template["id"],
        form_data=contract_form_data(),
    )

    response = client.post(
        f"/contracts/{contract['id']}/generate"
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        DOCX_MEDIA_TYPE
    )

    text = document_text(response.content)
    generated_document = Document(
        BytesIO(response.content)
    )

    assert "Д-101/26" in text
    assert "22 июля 2026" in text
    assert "Поставка оборудования" in text
    assert "1000.00 BYN" in text
    assert "ООО «Промас Инжиниринг»" in text
    assert "BY00TEST00000000000000000000" in text
    assert "ООО «Покупатель»" in text
    assert "BY11TEST11111111111111111111" in text
    assert "Петров П.П." in text
    assert "НЕ ДОЛЖНО ПОПАСТЬ" not in text
    assert get_template_variables(
        generated_document
    ) == set()

    contract_response = client.get(
        f"/contracts/{contract['id']}"
    )
    assert contract_response.status_code == 200
    generated_contract = contract_response.json()
    assert generated_contract["generated_file_name"] == (
        "Договор № Д-101_26.docx"
    )
    assert generated_contract[
        "generated_storage_path"
    ] is not None

    download_response = client.get(
        f"/contracts/{contract['id']}/download"
    )
    assert download_response.status_code == 200
    assert download_response.content == response.content

    events_response = client.get(
        f"/contracts/{contract['id']}/events"
    )
    assert events_response.status_code == 200
    events = events_response.json()
    assert events[0]["event_type"] == "generated"
    assert events[0]["event_data"]["template_id"] == (
        template["id"]
    )
    assert events[0]["event_data"][
        "version_number"
    ] == 1
    assert events[0]["event_data"]["file_sha256"] == (
        sha256(response.content).hexdigest()
    )

    versions_response = client.get(
        f"/contracts/{contract['id']}/versions"
    )
    assert versions_response.status_code == 200
    versions = versions_response.json()

    assert len(versions) == 1
    assert versions[0]["version_number"] == 1
    assert versions[0]["template_id"] == template["id"]
    assert versions[0]["template_name"] == (
        "Шаблон договора"
    )
    assert versions[0]["template_version"] == 1
    assert versions[0]["file_sha256"] == (
        sha256(response.content).hexdigest()
    )
    assert versions[0]["file_size_bytes"] == len(
        response.content
    )
    assert versions[0]["created_by_user_id"] is not None
    assert versions[0]["source_data"]["contract"][
        "number"
    ] == "Д-101/26"
    assert versions[0]["source_data"]["contract"][
        "amount"
    ] == "1000.00"
    assert versions[0]["source_data"]["counterparty"][
        "unp"
    ] == "900000101"
    assert versions[0]["source_data"]["organization"][
        "unp"
    ] == "590000001"

    version_download_response = client.get(
        f"/contracts/{contract['id']}/versions/1/download"
    )
    assert version_download_response.status_code == 200
    assert (
        version_download_response.content
        == response.content
    )


def test_generate_reports_all_missing_variables(
    client: TestClient,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    template_path = storage_root / "missing.docx"
    create_template_file(
        template_path,
        missing_variables_only=True,
    )
    template = upload_template(
        client,
        template_path,
        required_variables=[
            "organization.authority",
        ],
    )
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=template["id"],
    )

    response = client.post(
        f"/contracts/{contract['id']}/generate"
    )

    assert response.status_code == 422
    assert response.json() == {
        "detail": {
            "message": (
                "Не заполнены обязательные "
                "переменные шаблона"
            ),
            "missing_variables": [
                "contract.subject",
                "counterparty.account",
                "organization.authority",
            ],
        },
    }

    contract_response = client.get(
        f"/contracts/{contract['id']}"
    )
    assert contract_response.json()[
        "generated_storage_path"
    ] is None


def test_contract_rejects_wrong_template_type(
    client: TestClient,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    template_path = storage_root / "tz.docx"
    create_template_file(template_path)
    template = upload_template(
        client,
        template_path,
        template_type="technical_specification",
    )

    response = client.post(
        "/contracts",
        json={
            "counterparty_id": counterparty["id"],
            "template_id": template["id"],
            "number": "WRONG-TEMPLATE",
            "title": "Договор",
            "contract_date": "2026-07-22",
        },
    )

    assert response.status_code == 409
    assert response.json() == {
        "detail": (
            "Выбранный шаблон не предназначен "
            "для договоров"
        ),
    }


def test_contract_without_template_cannot_be_generated(
    client: TestClient,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=None,
    )

    response = client.post(
        f"/contracts/{contract['id']}/generate"
    )

    assert response.status_code == 409
    assert response.json() == {
        "detail": "Для договора не выбран шаблон",
    }


def test_archived_contract_cannot_be_generated(
    client: TestClient,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    template_path = storage_root / "source.docx"
    create_template_file(template_path)
    template = upload_template(
        client,
        template_path,
    )
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=template["id"],
        form_data=contract_form_data(),
    )
    archive_response = client.post(
        f"/contracts/{contract['id']}/archive"
    )
    assert archive_response.status_code == 200

    response = client.post(
        f"/contracts/{contract['id']}/generate"
    )

    assert response.status_code == 409


def test_repeated_generation_preserves_all_versions(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    update_organization_profile(client)
    template_path = storage_root / "source.docx"
    create_template_file(template_path)
    template = upload_template(
        client,
        template_path,
    )
    contract_data = create_contract(
        client,
        counterparty["id"],
        template_id=template["id"],
        form_data=contract_form_data(),
    )

    first_response = client.post(
        f"/contracts/{contract_data['id']}/generate"
    )
    assert first_response.status_code == 200

    db_session.expire_all()
    contract = db_session.get(
        Contract,
        contract_data["id"],
    )
    assert contract is not None
    first_path = Path(
        contract.generated_storage_path or ""
    )
    assert first_path.is_file()

    second_response = client.post(
        f"/contracts/{contract_data['id']}/generate"
    )
    assert second_response.status_code == 200

    db_session.expire_all()
    contract = db_session.get(
        Contract,
        contract_data["id"],
    )
    assert contract is not None
    second_path = Path(
        contract.generated_storage_path or ""
    )

    assert second_path.is_file()
    assert second_path != first_path
    assert first_path.is_file()

    versions_response = client.get(
        f"/contracts/{contract_data['id']}/versions"
    )
    assert versions_response.status_code == 200
    versions = versions_response.json()

    assert [
        version["version_number"]
        for version in versions
    ] == [2, 1]
    assert versions[0]["file_sha256"] == (
        sha256(second_response.content).hexdigest()
    )
    assert versions[1]["file_sha256"] == (
        sha256(first_response.content).hexdigest()
    )

    first_download_response = client.get(
        (
            f"/contracts/{contract_data['id']}"
            "/versions/1/download"
        )
    )
    assert first_download_response.status_code == 200
    assert (
        first_download_response.content
        == first_response.content
    )

    latest_download_response = client.get(
        f"/contracts/{contract_data['id']}/download"
    )
    assert latest_download_response.status_code == 200
    assert (
        latest_download_response.content
        == second_response.content
    )

    events_response = client.get(
        f"/contracts/{contract_data['id']}/events"
    )
    generated_events = [
        event
        for event in events_response.json()
        if event["event_type"] == "generated"
    ]
    assert len(generated_events) == 2
    assert generated_events[0]["event_data"][
        "version_number"
    ] == 2


def test_contract_update_preserves_generated_version(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    update_organization_profile(client)
    template_path = storage_root / "source.docx"
    create_template_file(template_path)
    template = upload_template(
        client,
        template_path,
    )
    contract_data = create_contract(
        client,
        counterparty["id"],
        template_id=template["id"],
        form_data=contract_form_data(),
    )
    generate_response = client.post(
        f"/contracts/{contract_data['id']}/generate"
    )
    assert generate_response.status_code == 200

    db_session.expire_all()
    contract = db_session.get(
        Contract,
        contract_data["id"],
    )
    assert contract is not None
    generated_path = Path(
        contract.generated_storage_path or ""
    )
    assert generated_path.is_file()

    update_response = client.patch(
        f"/contracts/{contract_data['id']}",
        json={
            "title": "Изменённый договор",
        },
    )

    assert update_response.status_code == 200
    assert update_response.json()[
        "generated_file_name"
    ] == "Договор № Д-101_26.docx"
    assert update_response.json()[
        "generated_storage_path"
    ] == str(generated_path)
    assert generated_path.is_file()

    download_response = client.get(
        f"/contracts/{contract_data['id']}/download"
    )
    assert download_response.status_code == 200
    assert (
        download_response.content
        == generate_response.content
    )

    second_generate_response = client.post(
        f"/contracts/{contract_data['id']}/generate"
    )
    assert second_generate_response.status_code == 200

    versions_response = client.get(
        f"/contracts/{contract_data['id']}/versions"
    )
    assert versions_response.status_code == 200
    versions = versions_response.json()

    assert [
        version["version_number"]
        for version in versions
    ] == [2, 1]
    assert versions[0]["source_data"]["contract"][
        "title"
    ] == "Изменённый договор"
    assert versions[1]["source_data"]["contract"][
        "title"
    ] == "Договор поставки"


def test_download_rejects_path_outside_contract_storage(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    update_organization_profile(client)
    template_path = storage_root / "source.docx"
    create_template_file(template_path)
    template = upload_template(
        client,
        template_path,
    )
    contract_data = create_contract(
        client,
        counterparty["id"],
        template_id=template["id"],
        form_data=contract_form_data(),
    )
    generate_response = client.post(
        f"/contracts/{contract_data['id']}/generate"
    )
    assert generate_response.status_code == 200

    outside_path = storage_root / "outside.docx"
    Document().save(outside_path)

    version = db_session.query(
        ContractDocumentVersion
    ).filter_by(
        contract_id=contract_data["id"],
        version_number=1,
    ).one()
    version.storage_path = str(outside_path)
    db_session.commit()

    response = client.get(
        (
            f"/contracts/{contract_data['id']}"
            "/versions/1/download"
        )
    )

    assert response.status_code == 404
    assert response.json() == {
        "detail": (
            "Сгенерированный DOCX-файл не найден"
        ),
    }


def test_missing_contract_document_version_returns_not_found(
    client: TestClient,
) -> None:
    counterparty = create_counterparty(client)
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=None,
    )

    response = client.get(
        f"/contracts/{contract['id']}/versions/1/download"
    )

    assert response.status_code == 404
    assert response.json() == {
        "detail": "Версия документа договора не найдена",
    }


def test_upload_contract_docx_creates_version(
    client: TestClient,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=None,
    )
    content = build_contract_document_content()

    response = client.post(
        (
            f"/contracts/{contract['id']}"
            "/versions/upload"
        ),
        files={
            "file": (
                "Договор контрагента.docx",
                content,
                DOCX_MEDIA_TYPE,
            ),
        },
    )

    assert response.status_code == 201
    version = response.json()
    assert version["version_number"] == 1
    assert version["source"] == "uploaded"
    assert version["template_id"] is None
    assert version["template_name"] is None
    assert version["template_version"] is None
    assert version["file_name"] == (
        "Договор контрагента.docx"
    )
    assert version["file_sha256"] == (
        sha256(content).hexdigest()
    )
    assert version["file_size_bytes"] == len(content)
    assert version["created_by_user_id"] is not None
    assert version["source_data"] == {
        "upload": {
            "original_file_name": (
                "Договор контрагента.docx"
            ),
            "content_type": DOCX_MEDIA_TYPE,
        },
    }

    stored_files = list(
        (
            storage_root
            / "generated"
            / "contracts"
        ).glob("*.docx")
    )
    assert len(stored_files) == 1
    assert stored_files[0].read_bytes() == content

    versions_response = client.get(
        f"/contracts/{contract['id']}/versions"
    )
    assert versions_response.status_code == 200
    assert versions_response.json() == [version]

    version_download_response = client.get(
        (
            f"/contracts/{contract['id']}"
            "/versions/1/download"
        )
    )
    assert version_download_response.status_code == 200
    assert version_download_response.content == content
    assert (
        version_download_response.headers[
            "content-disposition"
        ]
        .lower()
        .startswith("attachment;")
    )

    latest_download_response = client.get(
        f"/contracts/{contract['id']}/download"
    )
    assert latest_download_response.status_code == 200
    assert latest_download_response.content == content

    events_response = client.get(
        f"/contracts/{contract['id']}/events"
    )
    assert events_response.status_code == 200
    uploaded_event = events_response.json()[0]
    assert uploaded_event["event_type"] == "uploaded"
    assert uploaded_event["event_data"][
        "document_version_id"
    ] == version["id"]
    assert uploaded_event["event_data"][
        "version_number"
    ] == 1
    assert uploaded_event["event_data"][
        "file_name"
    ] == "Договор контрагента.docx"
    assert uploaded_event["event_data"][
        "file_sha256"
    ] == sha256(content).hexdigest()


def test_uploaded_and_generated_versions_share_sequence(
    client: TestClient,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    update_organization_profile(client)
    template_path = storage_root / "source.docx"
    create_template_file(template_path)
    template = upload_template(
        client,
        template_path,
    )
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=template["id"],
        form_data=contract_form_data(),
    )

    first_generated = client.post(
        f"/contracts/{contract['id']}/generate"
    )
    assert first_generated.status_code == 200

    uploaded_content = (
        build_contract_document_content(
            "Версия от контрагента"
        )
    )
    uploaded = client.post(
        (
            f"/contracts/{contract['id']}"
            "/versions/upload"
        ),
        files={
            "file": (
                "external.docx",
                uploaded_content,
                DOCX_MEDIA_TYPE,
            ),
        },
    )
    assert uploaded.status_code == 201
    assert uploaded.json()["version_number"] == 2

    second_generated = client.post(
        f"/contracts/{contract['id']}/generate"
    )
    assert second_generated.status_code == 200

    versions_response = client.get(
        f"/contracts/{contract['id']}/versions"
    )
    assert versions_response.status_code == 200
    versions = versions_response.json()

    assert [
        version["version_number"]
        for version in versions
    ] == [3, 2, 1]
    assert [
        version["source"]
        for version in versions
    ] == ["generated", "uploaded", "generated"]

    uploaded_download = client.get(
        (
            f"/contracts/{contract['id']}"
            "/versions/2/download"
        )
    )
    assert uploaded_download.status_code == 200
    assert uploaded_download.content == uploaded_content

    latest_download = client.get(
        f"/contracts/{contract['id']}/download"
    )
    assert latest_download.status_code == 200
    assert latest_download.content == (
        second_generated.content
    )


@pytest.mark.parametrize(
    ("file_name", "content"),
    [
        (
            "contract.txt",
            build_contract_document_content(),
        ),
        (
            "contract.docx",
            b"not a docx archive",
        ),
    ],
)
def test_upload_rejects_invalid_contract_document(
    client: TestClient,
    storage_root: Path,
    file_name: str,
    content: bytes,
) -> None:
    counterparty = create_counterparty(client)
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=None,
    )

    response = client.post(
        (
            f"/contracts/{contract['id']}"
            "/versions/upload"
        ),
        files={
            "file": (
                file_name,
                content,
                DOCX_MEDIA_TYPE,
            ),
        },
    )

    assert response.status_code == 422
    assert response.json() == {
        "detail": (
            "Необходимо загрузить корректный "
            "DOCX-файл"
        ),
    }
    assert client.get(
        f"/contracts/{contract['id']}/versions"
    ).json() == []
    assert not list(
        storage_root.rglob("*.docx")
    )


def test_upload_rejects_oversized_contract_document(
    client: TestClient,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=None,
    )
    content = b"x" * (
        contract_documents
        .MAX_CONTRACT_DOCUMENT_SIZE_BYTES
        + 1
    )

    response = client.post(
        (
            f"/contracts/{contract['id']}"
            "/versions/upload"
        ),
        files={
            "file": (
                "contract.docx",
                content,
                DOCX_MEDIA_TYPE,
            ),
        },
    )

    assert response.status_code == 413
    assert response.json() == {
        "detail": "Размер документа превышает 10 МБ",
    }
    assert client.get(
        f"/contracts/{contract['id']}/versions"
    ).json() == []
    assert not list(
        storage_root.rglob("*.docx")
    )


def test_archived_contract_rejects_document_upload(
    client: TestClient,
    storage_root: Path,
) -> None:
    counterparty = create_counterparty(client)
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=None,
    )
    archive_response = client.post(
        f"/contracts/{contract['id']}/archive"
    )
    assert archive_response.status_code == 200

    response = client.post(
        (
            f"/contracts/{contract['id']}"
            "/versions/upload"
        ),
        files={
            "file": (
                "contract.docx",
                build_contract_document_content(),
                DOCX_MEDIA_TYPE,
            ),
        },
    )

    assert response.status_code == 409
    assert response.json() == {
        "detail": (
            "Архивный договор сначала "
            "необходимо восстановить"
        ),
    }
    assert client.get(
        f"/contracts/{contract['id']}/versions"
    ).json() == []
    assert not list(
        storage_root.rglob("*.docx")
    )


def test_upload_removes_file_when_commit_fails(
    client: TestClient,
    db_session: Session,
    storage_root: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    counterparty = create_counterparty(client)
    contract = create_contract(
        client,
        counterparty["id"],
        template_id=None,
    )

    def fail_commit() -> None:
        raise RuntimeError("commit failed")

    with monkeypatch.context() as patch:
        patch.setattr(
            db_session,
            "commit",
            fail_commit,
        )

        with pytest.raises(
            RuntimeError,
            match="commit failed",
        ):
            client.post(
                (
                    f"/contracts/{contract['id']}"
                    "/versions/upload"
                ),
                files={
                    "file": (
                        "contract.docx",
                        build_contract_document_content(),
                        DOCX_MEDIA_TYPE,
                    ),
                },
            )

    assert not list(
        storage_root.rglob("*.docx")
    )
    assert (
        db_session.query(
            ContractDocumentVersion
        )
        .filter_by(contract_id=contract["id"])
        .count()
        == 0
    )
